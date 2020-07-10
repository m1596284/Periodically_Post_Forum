# -*- coding: UTF-8 -*-
import re
import os
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select
from selenium.webdriver.common.by import By
from selenium.webdriver import ActionChains 
from time import sleep
import docx
import win32com.client
import win32gui
import win32api
import win32con
import win32clipboard
import datetime
from urllib.parse import quote_plus as url_quoteplus
from urllib.parse import urlsplit
from urllib.request import urlretrieve
import urllib
from selenium.webdriver.common.by import By as WebBy
from selenium.webdriver.support.ui import Select as WebSelect
import pkg_resources.py2_warn
import numpy
import numpy.random.mtrand
import numpy.random.common
import numpy.random.bounded_integers
import numpy.random.entropy
from numpy.random import *
from PIL import ImageGrab
import pytesseract
from PIL import Image
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
pyPath, filename = os.path.split(__file__)
while 1:
    def readWord(ii):
        app=win32com.client.Dispatch('Word.Application')
        doc=app.Documents.Open(pyPath + "/" + Ac + "/檔案.docx")
        doc.Paragraphs(ii).Range.Copy()
        doc.Close()

    #貼上文章內容
    def okPaste():
        e_pasteword = driver.find_element_by_xpath('//*[@id="e_pasteword"]')
        sleep(0.5)
        e_pasteword.click()
        sleep(1)
        #tab
        win32api.keybd_event(0x09, 0, 0, 0)
        win32api.keybd_event(0x09, 0, win32con.KEYEVENTF_KEYUP, 0)
        sleep(0.5)
        #ctrl+V
        win32api.keybd_event(0x11, 0, 0, 0)
        win32api.keybd_event(0x56, 0, 0, 0)
        win32api.keybd_event(0x56, 0, win32con.KEYEVENTF_KEYUP, 0)
        win32api.keybd_event(0x11, 0, win32con.KEYEVENTF_KEYUP, 0)
        #paste
        e_pasteword_submit = driver.find_element_by_xpath('//*[@id="e_pasteword_submit"]')
        sleep(0.5)
        e_pasteword_submit.click()
        sleep(1)
        #enter
        win32api.keybd_event(0x0D, 0, 0, 0)
        win32api.keybd_event(0x0D, 0, win32con.KEYEVENTF_KEYUP, 0)
        sleep(0.5)

    #讀取並上傳圖片
    def onImage():
        # print(len(PictureList))
        e_image = driver.find_element_by_xpath('//*[@id="e_image"]')
        sleep(0.5)
        e_image.click()
        sleep(1)
        for picNum in range(len(PictureList)):
            # print(PictureList[picNum])
            #點擊上傳
            SWFUpload_0 = driver.find_element_by_xpath('//*[@id="SWFUpload_0"]')
            sleep(1)
            SWFUpload_0.click()
            sleep(5)
            # 點擊路徑
            for i2 in range (0,5):
                #tab 5
                win32api.keybd_event(0x09, 0, 0, 0)
                win32api.keybd_event(0x09, 0, win32con.KEYEVENTF_KEYUP, 0)
                sleep(0.5)
            #Enter
            win32api.keybd_event(0x0D, 0, 0, 0)
            win32api.keybd_event(0x0D, 0, win32con.KEYEVENTF_KEYUP, 0)
            sleep(1)
            win32clipboard.OpenClipboard()
            win32clipboard.EmptyClipboard()
            win32clipboard.SetClipboardText(pyPath+"/" + Ac)
            win32clipboard.CloseClipboard()
            win32api.keybd_event(0x11, 0, 0, 0)
            win32api.keybd_event(0x56, 0, 0, 0)
            win32api.keybd_event(0x56, 0, win32con.KEYEVENTF_KEYUP, 0)
            win32api.keybd_event(0x11, 0, win32con.KEYEVENTF_KEYUP, 0)
            sleep(1)
            win32api.keybd_event(0x0D, 0, 0, 0)
            win32api.keybd_event(0x0D, 0, win32con.KEYEVENTF_KEYUP, 0)
            sleep(1)
            # 選取輸入資料
            for i2 in range (0,5):
                #tab 5
                win32api.keybd_event(0x09, 0, 0, 0)
                win32api.keybd_event(0x09, 0, win32con.KEYEVENTF_KEYUP, 0)
                sleep(0.5)
            #輸入資料
            win32clipboard.OpenClipboard()
            win32clipboard.EmptyClipboard()
            win32clipboard.SetClipboardText(PictureList[picNum])
            win32clipboard.CloseClipboard()
            win32api.keybd_event(0x11, 0, 0, 0)
            win32api.keybd_event(0x56, 0, 0, 0)
            win32api.keybd_event(0x56, 0, win32con.KEYEVENTF_KEYUP, 0)
            win32api.keybd_event(0x11, 0, win32con.KEYEVENTF_KEYUP, 0)
            sleep(1)
            win32api.keybd_event(0x0D, 0, 0, 0)
            win32api.keybd_event(0x0D, 0, win32con.KEYEVENTF_KEYUP, 0)
            sleep(1)
        #按X關閉
        flbc = driver.find_element_by_xpath('//*[@class="flbc"]')
        sleep(0.5)
        flbc.click()
        sleep(1)
        #選取標題
        textSubject = driver.find_element_by_xpath('//*[@name="subject"]')
        sleep(0.5)
        textSubject.click()
        sleep(1)
        # 選取文章輸入格
        win32api.keybd_event(0x09, 0, 0, 0)
        win32api.keybd_event(0x09, 0, win32con.KEYEVENTF_KEYUP, 0)
        sleep(1)
        
    #按照片
    def clickPics(iii):
        e_image = driver.find_element_by_xpath('//*[@id="e_image"]')
        sleep(0.5)
        e_image.click()
        sleep(1)
        # picName = '//*[@id="
        # print(iii)
        # print(PictureList[iii])
        picClick = driver.find_element_by_xpath('//*[@title="%s"]' %PictureList[iii])
        sleep(0.5)
        picClick.click()
        sleep(1)
        #按X關閉
        flbc = driver.find_element_by_xpath('//*[@class="flbc"]')
        sleep(0.5)
        flbc.click()
        sleep(1)
    def allow_flash(driver, url):
        def _base_url(url):
            if url.find("://") == -1:
                url = "http://{}".format(url)
            urls = urlsplit(url)
            return "{}://{}".format(urls.scheme, urls.netloc)
        def _shadow_root(driver, element):
            return driver.execute_script("return arguments[0].shadowRoot", element)
        base_url = _base_url(url)
        driver.get("chrome://settings/content/siteDetails?site={}".format(url_quoteplus(base_url)))
        for iiii in range(20):
            # tab
            win32api.keybd_event(0x09, 0, 0, 0)
            win32api.keybd_event(0x09, 0, win32con.KEYEVENTF_KEYUP, 0)
            sleep(0.5)
        #enter
        win32api.keybd_event(0x0D, 0, 0, 0)
        win32api.keybd_event(0x0D, 0, win32con.KEYEVENTF_KEYUP, 0)
        sleep(2)        
        #up
        win32api.keybd_event(0x28, 0, 0, 0)
        win32api.keybd_event(0x28, 0, win32con.KEYEVENTF_KEYUP, 0)
        sleep(2)
        #enter
        win32api.keybd_event(0x0D, 0, 0, 0)
        win32api.keybd_event(0x0D, 0, win32con.KEYEVENTF_KEYUP, 0)
        sleep(2)

    #讀取密碼表
    f = open(pyPath + "/book.txt","r", encoding='UTF-8')
    readFile = f.readlines()
    f.close
    timePeriod = int(readFile[-1])
    # print(timePeriod)

    for acc in range(len(readFile)-1):
        try:
            # print(readFile[acc])
            Account = str(readFile[acc]).split('\t')
            Ac = Account[0]
            ID = Account[1]
            PW = Account[2].strip('\n')
            print(Ac)
            chrome_options = webdriver.ChromeOptions()
            chrome_options.add_argument('--allow-outdated-plugins')
            chrome_options.add_experimental_option("detach", True)
            driver = webdriver.Chrome(executable_path=pyPath + '/chromedriver.exe',chrome_options=chrome_options)
            driver.maximize_window()
            driver.get("https://www.chaforum.net/forum-171-1.html")
            sleep(5)
            #18 test
            try:
                periodaggre18 = driver.find_element_by_xpath('//*[@id="periodaggre18"]')
                sleep(1)
                periodaggre18.click()
                sleep(1)
                fwin_dialog_submit = driver.find_element_by_xpath('//*[@id="fwin_dialog_submit"]')
                sleep(1)
                fwin_dialog_submit.click()
                sleep(1)
            except:
                pass
            #Login
            username = driver.find_element_by_xpath('//*[@name="username"]')
            password = driver.find_element_by_xpath('//*[@name="password"]')
            BtnSure = driver.find_element_by_xpath('//*[@class="pn vm"]')
            username.send_keys(ID)
            password.send_keys(PW)
            sleep(1)
            BtnSure.click()
            sleep(3)
            allow_flash(driver,"https://www.chaforum.net/")
            driver.get("https://www.chaforum.net/forum-171-1.html")
            sleep(1)
            #開始新文章
            BtnNewPo = driver.find_element_by_xpath('//*[@id="newspecial"]')
            sleep(0.5)
            BtnNewPo.click()
            sleep(2)
            # #點擊驗證欄位
            #找地區
            textCity = driver.find_element_by_xpath('//*[@id="typeid_ctrl"]')
            sleep(0.5)
            textCity.click()
            sleep(1)
            #選地區
            file=docx.Document(pyPath + "/" + Ac + "/檔案.docx")
            City = file.paragraphs[0].text
            BtnCity = driver.find_element(By.XPATH, '//*[text()="%s"]' %City)
            sleep(0.5)
            BtnCity.click()
            sleep(1)
            #選取標題
            textSubject = driver.find_element_by_xpath('//*[@name="subject"]')
            sleep(0.5)
            textSubject.click()
            sleep(1)
            #讀取標題
            app=win32com.client.Dispatch('Word.Application')
            doc=app.Documents.Open(pyPath + "/" + Ac + "/檔案.docx")
            doc.Paragraphs(2).Range.Copy()
            doc.Close()
            #貼上標題
            win32api.keybd_event(0x11, 0, 0, 0)
            win32api.keybd_event(0x56, 0, 0, 0)
            win32api.keybd_event(0x56, 0, win32con.KEYEVENTF_KEYUP, 0)
            win32api.keybd_event(0x11, 0, win32con.KEYEVENTF_KEYUP, 0)
            sleep(1)
            # 選取文章輸入格
            win32api.keybd_event(0x09, 0, 0, 0)
            win32api.keybd_event(0x09, 0, win32con.KEYEVENTF_KEYUP, 0)
            sleep(1)
            #刪除原本文章
            win32api.keybd_event(0x11, 0, 0, 0)
            win32api.keybd_event(0x41, 0, 0, 0)
            win32api.keybd_event(0x41, 0, win32con.KEYEVENTF_KEYUP, 0)
            win32api.keybd_event(0x11, 0, win32con.KEYEVENTF_KEYUP, 0)
            sleep(1)
            win32api.keybd_event(0x2E, 0, 0, 0)
            win32api.keybd_event(0x2E, 0, win32con.KEYEVENTF_KEYUP, 0)
            sleep(1)
            #讀取圖片名冊
            PictureList =[]
            for dirPath, dirNames, fileNames in os.walk(pyPath+"/" + Ac):
                for f in fileNames:
                    if f.find("圖片")>-1:
                        # print(f)
                        PictureList.append(f)
            file=docx.Document(pyPath + "/" + Ac + "/檔案.docx")
            onImage()
            picCount = 0
            for i in range(3,len(file.paragraphs)+1):
                if file.paragraphs[i-1].text.find("圖片") > -1:
                    clickPics(picCount)
                    picCount = picCount+1
                    #Enter
                    win32api.keybd_event(0x0D, 0, 0, 0)
                    win32api.keybd_event(0x0D, 0, win32con.KEYEVENTF_KEYUP, 0)
                    sleep(1)
                else:
                    # print(i)
                    readWord(i)
                    okPaste()
            stone2 = 1
            while stone2 > 0:
                #處理驗證碼
                #換一個
                nextOne = driver.find_element_by_xpath("//a[contains(@onclick,'updateseccode')]")
                sleep(0.5)
                nextOne.click()
                #點擊驗證欄位
                seccodeverify = driver.find_element_by_xpath('//*[@name="seccodeverify"]')
                sleep(0.5)
                seccodeverify.click()
                #刪除原本文章
                win32api.keybd_event(0x11, 0, 0, 0)
                win32api.keybd_event(0x41, 0, 0, 0)
                win32api.keybd_event(0x41, 0, win32con.KEYEVENTF_KEYUP, 0)
                win32api.keybd_event(0x11, 0, win32con.KEYEVENTF_KEYUP, 0)
                sleep(0.5)
                win32api.keybd_event(0x2E, 0, 0, 0)
                win32api.keybd_event(0x2E, 0, win32con.KEYEVENTF_KEYUP, 0)
                sleep(0.5)
                #點擊驗證欄位
                seccodeverify = driver.find_element_by_xpath('//*[@name="seccodeverify"]')
                sleep(0.5)
                seccodeverify.click()
                #抓取驗證碼
                imgPic = driver.find_element_by_xpath("//img[contains(@onclick,'updateseccode')]")
                actionRight = ActionChains(driver)
                actionRight.context_click(imgPic).perform()
                sleep(0.5)
                win32api.keybd_event(0x59, 0, 0, 0)
                win32api.keybd_event(0x59, 0, win32con.KEYEVENTF_KEYUP, 0)
                sleep(0.5)
                img = ImageGrab.grabclipboard()
                img.save(pyPath + "/" + Ac + "/cha.png", 'PNG')
                img = Image.open(pyPath + "/" + Ac + "/cha.png").convert("L")
                pixdata = img.load()
                w = img.width
                h = img.height
                sto_pixel =img.getpixel((0,0))
                stone =0
                for i in range(w):
                    for ii in range(h):
                        cur_pixel = img.getpixel((i,ii))
                        # print(ii)
                        if (cur_pixel-sto_pixel)>10 or (cur_pixel-sto_pixel)<(-10):
                            key_pixel = cur_pixel
                            stone =1
                            break
                        else:
                            sto_pixel = cur_pixel
                        # pixelList.append(cur_pixel)
                        # print(cur_pixel)
                    if stone ==1:
                        break
                # print(key_pixel)
                for i in range(w):
                    for ii in range(h):
                        cur_pixel = img.getpixel((i,ii))
                        if cur_pixel == key_pixel:
                            pixdata[i,ii] = 0
                        else:
                            pixdata[i,ii] = 255
                # img.show()
                logAdd = pytesseract.image_to_string(img, lang='eng')
                logAdd = logAdd.replace(" ","")
                logAdd = logAdd.replace("°","V")
                logAdd = logAdd.replace("0","G")
                logAdd = logAdd.replace("I","J")
                logAdd = logAdd.replace("S","8")
                logAdd = logAdd.replace("Z","2")
                logAdd = logAdd.replace("1","Q")
                logAdd = logAdd.replace("g","9")
                logAdd = logAdd.replace("U","V")
                logAdd = logAdd.replace("¥","Y")
                logAdd = logAdd.replace("y","Y")
                logAdd = logAdd.replace("a","B")
                logAdd = logAdd.replace("o","C")
                logAdd = logAdd.replace("w","W")        
                logAdd = logAdd[0:4]
                seccodeverify = driver.find_element_by_xpath('//*[@name="seccodeverify"]')
                sleep(0.5)
                seccodeverify.send_keys(logAdd)
                sleep(0.5)
                postsubmit = driver.find_element_by_xpath('//*[@id="postsubmit"]')
                sleep(0.5)
                postsubmit.click()
                sleep(3)
                try:
                    e_fullswitcher = driver.find_element_by_xpath('//*[@id="e_fullswitcher"]')
                    sleep(0.5)
                except:
                    stone2 = 0
            driver.close()
        except:
            try:
                driver.close()
            except:
                pass
            continue

    print("循環完成: " + (datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
    print("下一個循環開始: " + (datetime.datetime.now()+datetime.timedelta(hours=timePeriod)).strftime("%Y-%m-%d %H:%M:%S"))
    print("等待" + str(timePeriod) + "小時")
    # timePeriod =0.001
    sleep(60*60*timePeriod)
    # sleep(timePeriod)



